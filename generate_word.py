import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = [
    {
        "Poin": "A",
        "Rekomendasi": "Autentikasi dan Manajemen Sesi",
        "Status": "Selesai",
        "Mengapa": "Jika server diretas dan database dicuri (data breach), hal pertama yang dicari peretas adalah tabel User dan Password untuk mengambil alih akun.",
        "Kenapa": "Kita menggunakan Bcrypt 12-rounds. Ini adalah algoritma enkripsi satu arah yang sangat lambat diproses oleh komputer. Jika peretas mencoba membobol jutaan password, Bcrypt 12-rounds akan membuat proses tersebut memakan waktu bertahun-tahun. Sesi (session) juga kita enkripsi agar peretas tidak bisa membajak sesi login orang lain dengan mencuri cookie di jaringan publik.",
        "Keterangan Implementasi": "Sistem menggunakan autentikasi berlapis:\n1. Password hashing menggunakan Bcrypt 12-round.\n2. Sesi berbasis database dengan enkripsi aktif (SESSION_ENCRYPT=true).\n3. Masa berlaku sesi dibatasi 300 menit.\n4. Session regeneration setelah login untuk mencegah Session Fixation.\n5. Session invalidation & regenerasi CSRF token saat logout.\n6. Kebijakan password kuat dengan indikator real-time.",
        "Bukti Kode": "BCRYPT_ROUNDS=12 | SESSION_DRIVER=database | SESSION_ENCRYPT=true | SESSION_LIFETIME=300\n$request->session()->regenerate()\n$request->session()->invalidate()\nPassword::min(8)->letters()->mixedCase()->numbers()->symbols()",
        "File Terkait": ".env (baris 16, 30-34)\nAuthenticatedSessionController.php (baris 46, 58-60)\nEditProfile.php",
        "Screenshot": "Halaman Login & Halaman Profil"
    },
    {
        "Poin": "B",
        "Rekomendasi": "Kontrol Akses",
        "Status": "Selesai",
        "Mengapa": "Ini merupakan mitigasi risiko reputasi (Reputation Risk). Jika akun staf regional dibajak, atau staf melakukan kesalahan ketik fatal, dan berita langsung tampil di publik mengatasnamakan instansi, hal itu akan menjadi masalah besar.",
        "Kenapa": "Role-Based Access Control (RBAC) memisahkan wewenang secara tegas. Data diisolasi agar staf daerah hanya bisa melihat dan mengedit datanya sendiri, sehingga kerusakan akibat compromised account (akun yang dibajak) sangat terlokalisir dan tidak menyebar ke seluruh sistem.",
        "Keterangan Implementasi": "Diimplementasikan menggunakan Spatie Laravel Permission dengan 2 peran (Admin dan Regional). Isolasi data dilakukan secara ketat di level query backend, di mana pengecekan peran dilakukan di sisi server via hasRole() yang tidak bisa dimanipulasi oleh klien.",
        "Bukti Kode": "use Spatie\\Permission\\Traits\\HasRoles\nclass User { use HasRoles; }\npublic function isAdmin(): bool { return $this->hasRole('admin'); }\nActivityResource: if ($user->hasRole('regional')) { $query->where('user_id', $user->id); }",
        "File Terkait": "User.php (baris 13, 18, 66-74)\nActivityResource.php (baris 478-488)",
        "Screenshot": "Daftar User (/admin/users) kolom Role"
    },
    {
        "Poin": "C",
        "Rekomendasi": "Validasi Input dan Pencegahan Injeksi",
        "Status": "Selesai",
        "Mengapa": "Serangan paling umum di internet dilakukan oleh robot otomatis (Bot). Bot akan mencoba memasukkan kode SQL jahat ke kolom input form untuk mengambil alih atau merusak isi database (SQL Injection).",
        "Kenapa": "Penggunaan antarmuka Eloquent ORM memastikan bahwa setiap input dari pengguna secara otomatis disanitasi dan dibersihkan (Prepared Statements). Selain itu, implementasi CAPTCHA memastikan bahwa form hanya bisa dikirim oleh manusia.",
        "Keterangan Implementasi": "Pencegahan injeksi data melalui Eloquent ORM tanpa adanya eksekusi raw SQL. Diterapkan pula validasi form sisi server yang ketat (required, maxLength, tipe data) dan CAPTCHA matematika pada login.",
        "Bukti Kode": "TextInput::make('extracted_title')->required()->maxLength(500)\nLoginRequest rules: 'email'=>['required','string','email'], 'captcha'=>['required','integer']",
        "File Terkait": "ActivityResource.php\nLoginRequest.php (baris 26-33)",
        "Screenshot": "Form Buat Kegiatan (/admin/activities/create)"
    },
    {
        "Poin": "D",
        "Rekomendasi": "Perlindungan Data dan Kriptografi",
        "Status": "Selesai",
        "Mengapa": "Sistem tidak boleh tanpa sengaja memuntahkan data sensitif ke publik. Data rahasia yang mengendap di dalam database maupun saat transit harus tersandi agar tidak bisa dibaca begitu saja.",
        "Kenapa": "Model User menggunakan fitur 'hashed casting' yang memastikan bahwa setiap kali ada password baru yang dimasukkan, sistem akan langsung mengenkripsinya secara sekejap sebelum data menyentuh database. Field sensitif juga di-hidden secara default dari keluaran API.",
        "Keterangan Implementasi": "Perlindungan data sensitif mencakup hashing password otomatis, penghapusan field password & remember_token dari serialisasi JSON ($hidden), sesi terenkripsi, dan kewajiban verifikasi password lama saat penggantian kata sandi.",
        "Bukti Kode": "protected $hidden = ['password','remember_token']\nprotected function casts(): array { return ['password' => 'hashed']; }\n->currentPassword()",
        "File Terkait": "User.php (baris 27-30, 32-38)\n.env\nEditProfile.php",
        "Screenshot": "Kode User.php & Halaman ganti profil"
    },
    {
        "Poin": "E",
        "Rekomendasi": "Penanganan Error dan Pencatatan Log",
        "Status": "Selesai",
        "Mengapa": "Peretas sering dengan sengaja memasukkan input aneh agar website menjadi Error. Jika konfigurasi salah, layar error akan membocorkan path folder server, versi database, dan celah teknis lainnya yang menjadi 'peta harta karun' bagi penyerang.",
        "Kenapa": "Kita mematikan debugger (APP_DEBUG=false) sehingga sistem hanya akan menampilkan halaman error standar (404/500) yang elegan kepada pengguna, sementara error aslinya disembunyikan dan dicatat (logged) secara rahasia di dalam server untuk ditinjau oleh teknisi.",
        "Keterangan Implementasi": "Halaman error kustom diimplementasikan tanpa menampilkan stack trace. Mode debug dimatikan. Log otomatis mencatat percobaan brute force (mencatat email, IP, user agent) ke dalam file harian terpusat (channel 'stack').",
        "Bukti Kode": "Log::warning('BRUTE FORCE DETECTED', ['email'=>$this->email, 'ip'=>$this->ip(), 'user_agent'=>$this->userAgent()])\nAPP_ENV=production\nAPP_DEBUG=false",
        "File Terkait": "LoginRequest.php (baris 96-101)\n.env\nresources/views/errors/layout.blade.php",
        "Screenshot": "Tampilan Error 404 (Akses URL sembarang)"
    },
    {
        "Poin": "F",
        "Rekomendasi": "Keamanan Komunikasi",
        "Status": "Selesai",
        "Mengapa": "Protokol HTTP biasa membuat data (seperti password dan dokumen yang diunggah) terbang melintasi jaringan internet (terutama WiFi publik) dalam bentuk teks telanjang yang bisa disadap dengan mudah (Man-in-the-Middle Attack).",
        "Kenapa": "Aplikasi dikonfigurasi untuk secara paksa menolak request HTTP biasa dan mengalihkannya ke jalur HTTPS. Dengan ini, semua komunikasi disandikan menggunakan sertifikat SSL/TLS sehingga data tidak bisa disadap di tengah jalan.",
        "Keterangan Implementasi": "Seluruh komunikasi data dijamin menggunakan HTTPS yang dipaksakan lewat URL::forceScheme('https') dan aturan server (.htaccess). Cookie sesi juga dikunci dengan flag Secure dan HttpOnly.",
        "Bukti Kode": "RewriteCond %{HTTPS} off\nRewriteRule ^(.*)$ https://%{HTTP_HOST}%{REQUEST_URI} [L,R=301]\nAppServiceProvider.php: URL::forceScheme('https')",
        "File Terkait": "public/.htaccess\nconfig/session.php\nAppServiceProvider.php",
        "Screenshot": "Potongan kode di AppServiceProvider.php (baris 27-29)"
    },
    {
        "Poin": "G",
        "Rekomendasi": "Perlindungan terhadap Serangan Umum",
        "Status": "Selesai",
        "Mengapa": "Serangan Brute Force dilakukan dengan mengirimkan ribuan percobaan login otomatis per menit. Jika tidak dibatasi, ini tidak hanya berisiko membobol password tetapi juga membuat server lumpuh (DDoS) karena kelebihan beban.",
        "Kenapa": "Mekanisme Rate Limit memastikan bahwa jika ada alamat IP yang mencoba dan gagal login sebanyak 3 kali beruntun, mereka akan 'dikunci' dan diblokir sementara. Waktu kuncian ini akan terus berlipat ganda hingga 15 menit jika penyerang terus memaksa.",
        "Keterangan Implementasi": "Perlindungan berlapis melalui CAPTCHA matematika, Rate Limiting progresif, CSRF token bawaan Laravel pada setiap form POST, pembatasan berdasar kombinasi IP & email, serta penonaktifan fitur 'Remember Me'.",
        "Bukti Kode": "RateLimiter::tooManyAttempts($this->throttleKey(), 3)\nProgressive: >=7=>900, >=4=>300, default=>60",
        "File Terkait": "AuthenticatedSessionController.php\nLoginRequest.php",
        "Screenshot": "Form Login (CAPTCHA) & Pesan Rate Limit"
    },
    {
        "Poin": "H",
        "Rekomendasi": "Keamanan API dan Web Service",
        "Status": "Selesai",
        "Mengapa": "Banyak sistem diretas karena membiarkan jalur komando (API endpoint) terbuka tanpa perlindungan, sehingga penyerang bisa langsung memanipulasi data tanpa harus melewati antarmuka website.",
        "Kenapa": "HAMDANS dibangun sebagai sistem monolith tertutup. Rute untuk mengubah atau menghapus data tidak pernah diekspos ke publik. Semua aksi penting harus dieksekusi dari dalam lorong panel admin yang dilindungi middleware sesi ketat.",
        "Keterangan Implementasi": "Rute publik hanya berjenis GET (baca data). Akses pengubahan data tersembunyi di balik rute '/admin' yang divalidasi oleh autentikasi cookie dan method canAccessPanel(). Percobaan akses tanpa login otomatis di-redirect.",
        "Bukti Kode": "Route publik (GET only): Route::get('/',...), Route::get('/kegiatan',...)\nTidak ada rute POST/PUT/DELETE untuk publik.",
        "File Terkait": "routes/web.php\nUser.php canAccessPanel()",
        "Screenshot": "Coba akses /admin tanpa login (Auto-redirect)"
    },
    {
        "Poin": "I",
        "Rekomendasi": "Pengelolaan File",
        "Status": "Selesai",
        "Mengapa": "Fitur unggah file adalah pintu masuk yang sangat empuk bagi peretas untuk menanamkan kode jahat (seperti virus atau script shell .php) langsung ke dalam jantung server.",
        "Kenapa": "Kita tidak hanya mengecek ekstensi file, melainkan memeriksa struktur DNA file (MIME type) di sisi server. Walaupun file dinamai 'gambar.jpg', jika isinya adalah script PHP, sistem akan menolak dan menghancurkannya. File yang lolos pun akan dikonversi dan dikompresi ulang secara paksa, yang otomatis menetralisir muatan virus.",
        "Keterangan Implementasi": "Menerapkan whitelisting format file (hanya JPEG, PNG, WebP) tervalidasi MIME. Isolasi penyimpanan file di disk 'public', batasan ukuran maksimal 5MB, serta mekanisme kompresi otomatis ke <100KB dan konversi ulang menjadi JPEG standar.",
        "Bukti Kode": "FileUpload::make('foto_dokumentasi')->image()->maxSize(5120)\ncompressImage(): memvalidasi $imageInfo['mime']",
        "File Terkait": "ActivityResource.php\nCreateActivity.php",
        "Screenshot": "Form Upload Kegiatan (Validasi ukuran & tipe)"
    },
    {
        "Poin": "J",
        "Rekomendasi": "Pengendalian Kode dan Dependensi",
        "Status": "Selesai",
        "Mengapa": "Menggunakan bahasa pemrograman atau modul usang sama seperti membiarkan pintu rumah rusak dan tidak dikunci, karena daftar celah keamanannya sudah tersebar luas di internet dan diketahui semua peretas.",
        "Kenapa": "Kita memastikan fondasi aplikasi berdiri di atas framework paling stabil dan mutakhir (Laravel 12). Ketergantungan eksternal (dependencies) dikunci menggunakan composer.lock untuk mencegah masuknya versi pustaka jahat atau belum teruji saat deploy.",
        "Keterangan Implementasi": "Digunakan framework PHP mutakhir dengan dukungan perbaikan keamanan jangka panjang. Penggunaan modul eksternal dikelola ketat via Composer dan versi eksaknya dikunci secara absolut.",
        "Bukti Kode": "composer.json: laravel/framework ^12.0, filament/filament ^3.0, spatie/laravel-permission ^7.1",
        "File Terkait": "composer.json",
        "Screenshot": "Isi file composer.json bagian 'require'"
    },
    {
        "Poin": "K",
        "Rekomendasi": "Keamanan Konfigurasi Sistem",
        "Status": "Selesai",
        "Mengapa": "Pertahanan tingkat militer sekalipun akan runtuh jika administrator utamanya menggunakan password default seperti '123456' atau bisa diganti oleh sembarang orang di komputernya.",
        "Kenapa": "Sistem memaksa pemberlakuan kebijakan sandi rumit (harus ada huruf, angka, simbol). Selain itu, untuk melakukan penggantian password rahasia, pengguna diwajibkan mengetikkan ulang password lama mereka untuk memastikan bahwa itu memang pemilik asli akun, bukan orang yang menemukan laptop terbuka.",
        "Keterangan Implementasi": "Sistem mematikan debugger, mewajibkan password kuat (kombinasi 4 komponen), memverifikasi password lama sebelum menyimpan perubahan, serta memastikan lingkungan berjalan dalam mode APP_ENV=production.",
        "Bukti Kode": "APP_ENV=production, APP_DEBUG=false\nPassword::min(8)->letters()->mixedCase()->numbers()->symbols()\n->currentPassword()",
        "File Terkait": ".env (produksi)\nEditProfile.php",
        "Screenshot": "Halaman Profil (Verifikasi Password Saat Ini)"
    },
    {
        "Poin": "L",
        "Rekomendasi": "Pengamanan Logika Bisnis",
        "Status": "Selesai",
        "Mengapa": "Sistem harus dirancang pesimis untuk mencegah terjadinya *human error* atau usaha *by-pass* (mengakali sistem), misalnya saat staf daerah memaksa agar berita mereka langsung terpublikasi.",
        "Kenapa": "Penerapan logika 'Force Pending' di level terdalam aplikasi memastikan bahwa perintah apapun dari klien akan ditimpa menjadi status 'Pending'. Bahkan, jika berita berstatus 'Approved' diam-diam disunting lagi oleh staf daerah, sistem akan seketika mencabut status tayangnya dan mengembalikannya ke 'Pending' untuk dicek ulang oleh pusat.",
        "Keterangan Implementasi": "Staf regional TIDAK bisa menerbitkan berita secara langsung. Hak persetujuan (Approve) eksklusif dimiliki peran Admin. Aktivitas sunting (edit) pada dokumen yang telah disetujui akan otomatis merevokasi status publikasi kembali menjadi Pending tanpa menghapus tanggal rilis awal.",
        "Bukti Kode": "CreateActivity: $data['status'] = ActivityStatus::Pending\nEditActivity: if (Approved) { $data['status'] = Pending; unset($data['approved_at']); }",
        "File Terkait": "CreateActivity.php\nEditActivity.php",
        "Screenshot": "Form Create Kegiatan (tanpa Status) & Daftar Kegiatan (Pending)"
    }
]

# Membuat Dokumen Word
doc = Document()

# Judul Dokumen
title = doc.add_heading('Laporan Detail Penyelesaian Rekomendasi Keamanan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Dokumen ini berisi rincian teknis beserta penjelasan rasionalisasi (Kenapa dan Mengapa) penyelesaian untuk 12 poin rekomendasi keamanan aplikasi portal HAMDANS.')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

for item in data:
    # Heading per poin
    heading = doc.add_heading(f"Poin {item['Poin']}. {item['Rekomendasi']}", level=1)
    
    # Status
    p_status = doc.add_paragraph()
    p_status.add_run('Status Penyelesaian: ').bold = True
    run_status = p_status.add_run(item['Status'])
    run_status.font.color.rgb = RGBColor(0, 128, 0) # Hijau
    
    # Kenapa dan Mengapa
    doc.add_heading('Analisis Risiko & Rasionalisasi', level=2)
    p_mengapa = doc.add_paragraph()
    p_mengapa.add_run('Mengapa Diperlukan?\n').bold = True
    p_mengapa.add_run(item['Mengapa'])
    
    p_kenapa = doc.add_paragraph()
    p_kenapa.add_run('Kenapa Implementasinya Demikian?\n').bold = True
    p_kenapa.add_run(item['Kenapa'])
    
    # Keterangan
    doc.add_heading('Detail Implementasi', level=2)
    doc.add_paragraph(item['Keterangan Implementasi'])
    
    # Bukti Kode
    doc.add_heading('Bukti / Potongan Kode', level=2)
    p_code = doc.add_paragraph(item['Bukti Kode'])
    for run in p_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
    # File Terkait
    doc.add_heading('Lokasi File Terkait', level=2)
    doc.add_paragraph(item['File Terkait'])
    
    # Screenshot
    doc.add_heading('Instruksi Screenshot / Bukti Visual', level=2)
    p_ss = doc.add_paragraph()
    p_ss.add_run('▶ ').bold = True
    p_ss.add_run(item['Screenshot']).italic = True
    
    doc.add_paragraph('_' * 60) # Garis pemisah

# Summary Quote at the end
doc.add_paragraph()
quote_box = doc.add_paragraph()
quote_box.add_run('Ringkasan Eksekutif untuk Auditor:\n').bold = True
quote_run = quote_box.add_run('“Sistem HAMDANS tidak hanya mengandalkan keamanan di tampilan depan (Frontend), tetapi membangun dinding pertahanan dari dalam mesinnya (Backend). Meskipun seorang peretas berhasil melewati satu lapis pengamanan, sistem akan mendeteksi kejanggalan dan memblokirnya di lapisan berikutnya. Sistem ini juga mencegah terjadinya kesalahan manusia (Human Error) dari staf daerah yang dapat berdampak pada kredibilitas institusi publik.”')
quote_run.italic = True
quote_box.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Simpan Dokumen
file_name = 'Laporan_Penyelesaian_Keamanan.docx'
doc.save(file_name)
print(f"✅ Dokumen Word berhasil dibuat: {file_name}")
